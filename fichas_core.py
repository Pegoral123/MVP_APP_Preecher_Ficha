"""
fichas_core.py
------------------------------------------------------------------
Lógica central do Gerador de Fichas de Matrícula (UNIFESO).
Usada tanto pela versão CLI (gerar_fichas_cli_refatorado.py) quanto pela
versão com interface gráfica (gerar_fichas_gui_refatorado.py).

Conversão docx -> pdf:
  1. Tenta Microsoft Word (docx2pdf).
  2. Se falhar, tenta LibreOffice (soffice --headless).
  3. Se ambos falharem, erro claro para o usuário.
"""
import re
import os
import shutil
import subprocess
from pathlib import Path

import pdfplumber
from docxtpl import DocxTemplate
from pypdf import PdfWriter


# ---------------------------------------------------------------------------
# Exceção própria
# ---------------------------------------------------------------------------
class FichasError(Exception):
    """Erro esperado do processo (para mostrar mensagem amigável na GUI)."""
    pass


# ---------------------------------------------------------------------------
# Validações de integridade (pré-processamento)
# ---------------------------------------------------------------------------
def _validar_template(template_path: str):
    """Confirma que o .docx contém os placeholders obrigatórios."""
    from docx import Document
    texto_completo = []
    doc = Document(template_path)
    for p in doc.paragraphs:
        texto_completo.append(p.text)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                texto_completo.append(celula.text)

    conteudo = " ".join(texto_completo)
    faltantes = []
    if "{{ nome }}" not in conteudo and "{{nome}}" not in conteudo:
        faltantes.append("{{ nome }}")
    if ("{{ matricula }}" not in conteudo
            and "{{matricula}}" not in conteudo):
        faltantes.append("{{ matricula }}")

    if faltantes:
        raise FichasError(
            f"Template inválido: faltam os placeholders "
            f"{', '.join(faltantes)} no arquivo '{template_path}'.\n"
            "O template deve conter {{ nome }} e {{ matricula }} "
            "nos locais onde os dados devem ser preenchidos."
        )


def _validar_permissoes_saida(saida_pdf: str):
    """Confirma que a pasta de destino tem permissão de escrita."""
    pasta = Path(saida_pdf).parent
    if pasta.exists():
        if not os.access(str(pasta), os.W_OK):
            raise FichasError(
                f"Sem permissão de escrita na pasta de destino:\n{pasta}"
            )
    else:
        # Tenta criar a pasta pra testar permissão
        try:
            pasta.mkdir(parents=True, exist_ok=True)
        except Exception as e:
            raise FichasError(
                f"Não foi possível criar a pasta de destino:\n{pasta}\nDetalhe: {e}"
            )


# ---------------------------------------------------------------------------
# Detecção de Word instalado (via win32com)
# ---------------------------------------------------------------------------
def _word_instalado():
    """Retorna True se o Microsoft Word está instalado e funcional."""
    try:
        import win32com.client
        # Tenta múltiplos ProgIDs (versões diferentes do Word)
        for prog_id in ["Word.Application", "Word.Application.16", "Word.Application.15"]:
            try:
                word = win32com.client.Dispatch(prog_id)
                word.Quit()
                return True
            except Exception:
                continue
        return False
    except ImportError:
        return False


def _libreoffice_instalado():
    """Retorna True se o LibreOffice está acessível no PATH ou em caminhos
    comuns de instalação no Windows."""
    # 1. Tenta no PATH
    if shutil.which("soffice"):
        return True

    # 2. Caminhos comuns do LibreOffice no Windows
    caminhos_comuns = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for c in caminhos_comuns:
        if Path(c).exists():
            return True
    return False


def _caminho_libreoffice():
    """Retorna o caminho completo do executável soffice.exe ou None."""
    if shutil.which("soffice"):
        return shutil.which("soffice")
    caminhos_comuns = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for c in caminhos_comuns:
        if Path(c).exists():
            return c
    return None


# ---------------------------------------------------------------------------
# Extração de alunos da ata PDF
# ---------------------------------------------------------------------------
def extrair_alunos(pdf_path: str, log=print):
    """Extrai (RA, Nome) de todas as páginas da ata.

    O PDF do TOTVS (Alunos por Período) coloca os dados no formato:
        <RA> <NOME COMPLETO> MATRICULADO CIÊNCIA DA
    Ex: "06019004 ABNER CAVALCANTI ATOUGUIA MATRICULADO CIÊNCIA DA"

    O RA são dígitos no início da linha, o nome é tudo entre o RA
    e a palavra MATRICULADO.
    """
    alunos = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            for line in text.split("\n"):
                linha = line.strip()
                # Formato: RA NOME MATRICULADO ...
                m = re.match(
                    r"^(\d{6,10})\s+(.+?)\s+MATRICULADO", linha
                )
                if m:
                    ra = m.group(1)
                    nome = m.group(2).strip()
                    alunos.append({"ra": ra, "nome": nome})

    if not alunos:
        raise FichasError(
            "Nenhum aluno MATRICULADO encontrado na ata.\n\n"
            "Possíveis causas:\n"
            "  • O PDF não é o relatório 'Alunos por Período' "
            "exportado do TOTVS.\n"
            "  • O layout do relatório mudou e o regex de extração "
            "precisa ser atualizado.\n\n"
            "Formato esperado em cada linha:\n"
            "  <RA> <NOME COMPLETO> MATRICULADO ..."
        )

    alunos.sort(key=lambda a: a["nome"])
    log(f"      {len(alunos)} alunos encontrados (status MATRICULADO).")
    return alunos


# ---------------------------------------------------------------------------
# Geração das fichas .docx
# ---------------------------------------------------------------------------
def gerar_fichas_docx(alunos, template_path: str, out_dir: Path,
                      log=print, progress_callback=None):
    """Preenche o template para cada aluno, salvando os .docx individuais.

    Args:
        progress_callback: opcional, chamado com (atual, total) a cada ficha.
    """
    out_dir.mkdir(parents=True, exist_ok=True)
    caminhos = []
    total = len(alunos)
    for i, aluno in enumerate(alunos, start=1):
        doc = DocxTemplate(template_path)
        doc.render({"nome": aluno["nome"], "matricula": aluno["ra"]})
        out_path = out_dir / f"ficha_{aluno['ra']}.docx"
        doc.save(out_path)
        caminhos.append(out_path)

        if progress_callback:
            progress_callback(i, total)

        if i % 25 == 0 or i == total:
            log(f"      {i}/{total} fichas .docx geradas...")
    return caminhos


# ---------------------------------------------------------------------------
# Conversão docx -> PDF com fallback (Word → LibreOffice)
# ---------------------------------------------------------------------------
def _converter_via_word(out_dir: Path, log=print):
    """Converte todos os .docx da pasta usando Microsoft Word (docx2pdf)."""
    from docx2pdf import convert
    log("      Tentando converter via Microsoft Word...")
    convert(str(out_dir))
    log("      Conversão via Word concluída.")


def _converter_via_libreoffice(out_dir: Path, log=print):
    """Converte todos os .docx da pasta usando LibreOffice headless."""
    soffice = _caminho_libreoffice()
    if not soffice:
        raise FichasError("LibreOffice não encontrado no sistema.")

    log(f"      Tentando converter via LibreOffice ({soffice})...")

    cmd = [
        soffice,
        "--headless",
        "--convert-to", "pdf",
        "--outdir", str(out_dir),
    ]
    docx_files = list(out_dir.glob("*.docx"))
    if not docx_files:
        raise FichasError("Nenhum arquivo .docx encontrado para converter.")
    cmd.extend([str(f) for f in docx_files])

    resultado = subprocess.run(
        cmd,
        capture_output=True,
        text=True,
        timeout=300,
        cwd=str(out_dir),
    )

    if resultado.returncode != 0:
        raise FichasError(
            f"LibreOffice falhou na conversão.\n"
            f"stdout: {resultado.stdout}\n"
            f"stderr: {resultado.stderr}"
        )

    log("      Conversão via LibreOffice concluída.")


def converter_para_pdf_em_lote(docx_paths, out_dir: Path, log=print):
    """Converte todos os .docx da pasta para PDF.

    Tenta primeiro Microsoft Word (docx2pdf).
    Se falhar, tenta LibreOffice (soffice --headless).
    Se ambos falharem, levanta FichasError com mensagem explicativa.
    """
    erros = []

    # --- Tentativa 1: Microsoft Word ---
    if _word_instalado():
        try:
            _converter_via_word(out_dir, log)
            return [out_dir / (p.stem + ".pdf") for p in docx_paths]
        except Exception as e:
            erros.append(f"Microsoft Word: {e}")
            log(f"      Word falhou: {e}")
    else:
        erros.append("Microsoft Word não está instalado.")

    # --- Tentativa 2: LibreOffice ---
    if _libreoffice_instalado():
        try:
            _converter_via_libreoffice(out_dir, log)
            return [out_dir / (p.stem + ".pdf") for p in docx_paths]
        except Exception as e:
            erros.append(f"LibreOffice: {e}")
            log(f"      LibreOffice falhou: {e}")
    else:
        erros.append("LibreOffice não está instalado ou não foi encontrado no PATH.")

    # --- Nenhum funcionou ---
    msg = "Falha na conversão para PDF. Nenhum conversor disponível:\n\n"
    for erro in erros:
        msg += f"  • {erro}\n"
    msg += (
        "\nPara resolver, instale uma das opções:\n"
        "  • Microsoft Word (recomendado)\n"
        "  • LibreOffice (gratuito: https://pt-br.libreoffice.org/)\n"
    )
    raise FichasError(msg)


# ---------------------------------------------------------------------------
# Junção dos PDFs
# ---------------------------------------------------------------------------
def juntar_pdfs(pdf_paths, saida: str, log=print):
    writer = PdfWriter()
    faltando = []
    for p in pdf_paths:
        if p.exists():
            writer.append(str(p))
        else:
            faltando.append(p.name)
    if faltando:
        log(f"      Aviso: {len(faltando)} ficha(s) não foram convertidas "
            f"e ficaram de fora.")
    with open(saida, "wb") as f:
        writer.write(f)


# ---------------------------------------------------------------------------
# Pipeline completo
# ---------------------------------------------------------------------------
def gerar_fichas_completo(ata_pdf: str, template_docx: str, saida_pdf: str,
                          work_dir: Path = None, log=print, limpar_temp=True,
                          progress_callback=None):
    """
    Roda o pipeline inteiro: extrair alunos -> gerar docx -> converter pdf -> juntar.
    Retorna a quantidade de alunos processados.

    Args:
        progress_callback: opcional, chamado com (etapa, atual, total)
                           etapa: "docx" ou "pdf"
    """
    if work_dir is None:
        work_dir = Path(saida_pdf).parent / "_fichas_temp"

    # --- Validações prévias ---
    log("0/5 - Validando arquivos de entrada...")

    if not Path(ata_pdf).exists():
        raise FichasError(f"Ata PDF não encontrada:\n{ata_pdf}")
    if not Path(template_docx).exists():
        raise FichasError(f"Template .docx não encontrado:\n{template_docx}")

    _validar_template(template_docx)
    _validar_permissoes_saida(saida_pdf)
    log("      Validações OK.")

    # --- Etapa 1: Extrair alunos ---
    log("1/5 - Lendo ata e extraindo alunos...")
    alunos = extrair_alunos(ata_pdf, log=log)

    # --- Etapa 2: Gerar .docx ---
    log("2/5 - Gerando fichas .docx individuais...")
    if progress_callback:
        def _cb_docx(atual, total):
            progress_callback("docx", atual, total)
    else:
        _cb_docx = None
    docx_paths = gerar_fichas_docx(alunos, template_docx, work_dir,
                                   log=log, progress_callback=_cb_docx)

    # --- Etapa 3: Verificar conversores ANTES de tentar ---
    log("3/5 - Verificando conversores disponíveis...")
    tem_word = _word_instalado()
    tem_libre = _libreoffice_instalado()

    if tem_word:
        log("      Microsoft Word: OK")
    else:
        log("      Microsoft Word: NÃO ENCONTRADO")

    if tem_libre:
        log(f"      LibreOffice: OK ({_caminho_libreoffice()})")
    else:
        log("      LibreOffice: NÃO ENCONTRADO")

    if not tem_word and not tem_libre:
        raise FichasError(
            "Nenhum conversor PDF encontrado no sistema.\n\n"
            "Para gerar os PDFs, é necessário ter instalado:\n"
            "  • Microsoft Word (recomendado)\n"
            "  • LibreOffice (gratuito: https://pt-br.libreoffice.org/)\n\n"
            "Instale um deles e tente novamente."
        )

    # --- Etapa 4: Converter para PDF ---
    log("4/5 - Convertendo fichas para PDF...")
    if progress_callback:
        progress_callback("pdf", 0, len(alunos))
    pdf_paths = converter_para_pdf_em_lote(docx_paths, work_dir, log=log)
    if progress_callback:
        progress_callback("pdf", len(alunos), len(alunos))

    # --- Etapa 5: Juntar PDFs ---
    log("5/5 - Juntando tudo em um PDF único para impressão...")
    juntar_pdfs(pdf_paths, saida_pdf, log=log)

    if limpar_temp:
        shutil.rmtree(work_dir, ignore_errors=True)

    log(f"\nPronto! {len(alunos)} fichas geradas em: {saida_pdf}")
    return len(alunos)